"""

 Created on 2025/1/9  
 @author: yks 
 Path: D:/yhq/python31105/langchain.py
 use:documents&&floorspace
"""
# -*- coding: utf-8 -*-
import os
from fastapi import FastAPI
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_community.chat_models import QianfanChatEndpoint
from langchain_community.embeddings import QianfanEmbeddingsEndpoint
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_core.language_models.chat_models import HumanMessage
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnableWithMessageHistory, RunnableLambda, RunnablePassthrough
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langserve import add_routes
import os
from docx import Document as doc

qianfan_ak = os.getenv("QIANFAN_AK")
qianfan_sk = os.getenv("QIANFAN_SK")
# print(qianfan_sk,qianfan_ak)
# 得到模型：千帆
qianfan_chat = QianfanChatEndpoint(
    model="ERNIE-3.5-8K",
    temperature=0.2,
    timeout=30,
    api_key=qianfan_ak,
    secret_key=qianfan_sk,

)
doc = doc('just.docx')
texts = []
#遍历文档中的每一个段落并提取文本
for para in doc.paragraphs:
    texts.append(para.text)
print(len(''.join(texts)))
# 准备测试数据
# documents = [
#     # Document(
#     #     page_content='节约支出',
#     #     metadata={"source": 'good'}
#     # ),
#     # Document(
#     #     page_content='增加收入',
#     #     metadata={"source": 'good'}
#     # ),
#     # Document(
#     #     page_content='裁减人员',
#     #     metadata={"source": 'shabi'}
#     # ),
#     # Document(
#     #     page_content='关闭公司',
#     #     metadata={"source": 'notgood'}
#     # ),
#     # Document(
#     #     page_content='降低成本，增加效率',
#     #     metadata={"source": 'good'}
#     # ),
#     # Document(
#     #     page_content='降低效率，增加成本',
#     #     metadata={"source": 'shabi'}
#     # ),
#
# ]
documents = [
Document(
    page_content=''.join(texts),
    metadata={"source": 'shabi'}
)
]
# 实例化一个向量空间

splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=30)
res = splitter.split_documents(documents)
vector_store = Chroma.from_documents(res, embedding=QianfanEmbeddingsEndpoint())
retriver = RunnableLambda(vector_store.from_documents())
# print(retriver.batch(['降本增效','开除 ']))

#构造解析器
param = StrOutputParser()
#提示模板
message='''
使用提供的上下文回答这个问题.
{question}
上下文：
{context}
'''
prompt_template = ChatPromptTemplate.from_messages([
    ('human',message)
])

#RunnablePassthrough允许问题之后再传递
chain={'question':RunnablePassthrough(), 'context' :retriver}  | prompt_template | qianfan_chat|param
res=chain.invoke('简要说明这篇文章的内容')
print(res)