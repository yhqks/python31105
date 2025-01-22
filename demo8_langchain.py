import os
from langchain.chains.retrieval import create_retrieval_chain
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import QianfanEmbeddingsEndpoint
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnableWithMessageHistory
from langchain_community.chat_models import QianfanChatEndpoint
from langchain_core.output_parsers import StrOutputParser
from langchain_community.chat_message_histories import ChatMessageHistory
import docx
from langchain_core.documents import Document
# Step 1: Data Collection
file_path = 'C:\Users\yks\Desktop\AIOPS资料及代码\AIOPS资料及代码\资料\基站全生命周期智能运维系统-结题汇报20240318.pptx'  # Replace with your local file path

def load_docx_file(file_path):# 读取.docx文件
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

docs=load_docx_file(file_path)
#形成为文档格式
documents = [
Document(
    page_content=docs,
    metadata={"source": 'text'}
)
]

# Step 2: Data Processing 切割文档内容
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
res = splitter.split_documents(documents)

# Step 3: Vectorization 向量化
qianfan_ak = os.getenv("QIANFAN_AK")
qianfan_sk = os.getenv("QIANFAN_SK")
embedding_model = QianfanEmbeddingsEndpoint(api_key=qianfan_ak, secret_key=qianfan_sk)
# 创建向量存储 persist_directory='chroma_data_dir' 保存向量数据
vectorstore = Chroma.from_documents(documents=res, embedding=embedding_model)

# Step 4: Storage 创建搜索器
retriever = vectorstore.as_retriever()

# Step 5: Retrieval 
model = QianfanChatEndpoint(
    model="ERNIE-4.0-8K",
    temperature=0.2,
    timeout=180,
    api_key=qianfan_ak,
    secret_key=qianfan_sk,
)

message = '''
你是一个助手，可以来回答关于文章内的问题,\n
{context}
'''
# 创建提示
prompt_template = ChatPromptTemplate.from_messages([
    ('human', message),
    MessagesPlaceholder('chat_history'),
    ('human', "{input}")
])
#回答解析器
param = StrOutputParser()
# 创建检索链
model = prompt_template|model | param
chain = create_retrieval_chain(retriever, model)
# Step 6: Chat 保存历史数据
store = {}
# 获取历史记录
def get_session_history(session_id: str):
    if session_id not in store:
        store[session_id] = ChatMessageHistory()
    return store[session_id]
# 创建链
result_chain = RunnableWithMessageHistory(
    chain,
    get_session_history,
    input_messages_key='input',
    history_messages_key='chat_history',
    output_messages_key='/'
)

# Example query
input_text = "这个文档的主要内容是什么？"
resp = result_chain.invoke(
    {'input': input_text},
    config={'configurable': {'session_id': 'yks'}}
)
print(resp['answer'])